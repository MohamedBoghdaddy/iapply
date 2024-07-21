# Import necessary libraries and modules
import textract
from docx import Document
from pdfminer.high_level import extract_text as extract_pdf_text
from pdfminer.pdfparser import PDFSyntaxError
import natural
from sklearn.cluster import KMeans
import requests
from datetime import datetime
from dateutil.parser import parse
from marshmallow import ValidationError
import smtplib
import ssl
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import six
from dotenv import load_dotenv
import os

# Load environment variables from .env file
load_dotenv()

env_path = os.getenv('../server/.env')
if env_path:
    load_dotenv(dotenv_path=env_path)

# Constants
ANTHROPICS_API_KEY = os.getenv('ANTHROPICS_API_KEY')

# CV Text Extraction
def parse_pdf(file_path):
    """Extract text from a PDF file."""
    try:
        txt = extract_pdf_text(file_path).strip()
        if not txt:
            raise ValueError("Empty PDF content")
        return txt
    except FileNotFoundError:
        print("File not found or path is incorrect")
        raise
    except PDFSyntaxError:
        print("Not a valid PDF file")
        raise

def parse_docx(file_path):
    """Extract text from a DOCX file."""
    try:
        doc = Document(file_path)
        txt = "\n".join([paragraph.text for paragraph in doc.paragraphs]).strip()
        if not txt:
            raise ValueError("Empty DOCX content")
        return txt
    except FileNotFoundError:
        print("File not found or path is incorrect")
        raise
    except Exception as e:
        print("Error parsing DOCX file:", e)
        raise

def parse_doc(file_path):
    """Extract text from a DOC file."""
    try:
        txt = textract.process(file_path).decode("utf-8").strip()
        if not txt:
            raise ValueError("Empty DOC content")
        return txt
    except FileNotFoundError:
        print("File not found or path is incorrect")
        raise
    except Exception as e:
        print("Error parsing DOC file:", e)
        raise

def extract_text(file_path):
    """Extract text from a file based on its extension."""
    if file_path.endswith(".pdf"):
        return parse_pdf(file_path)
    elif file_path.endswith(".docx"):
        return parse_docx(file_path)
    elif file_path.endswith(".doc"):
        return parse_doc(file_path)
    else:
        raise ValueError("Unsupported file format")

# CV Analyzer
class CVAnalyzer:
    def __init__(self, cv_text):
        self.cv_text = cv_text
        self.tokenizer = natural.WordTokenizer()
        self.tfidf = natural.TfIdf()

    def analyze_skills(self):
        """Analyze skills from CV using TF-IDF"""
        words = self.tokenizer.tokenize(self.cv_text)
        self.tfidf.addDocument(words)
        return self.tfidf

# Job Clustering
def cluster_jobs(jobs, num_clusters=5):
    """Cluster job descriptions using K-Means"""
    descriptions = [job['description'] for job in jobs]
    tfidf = natural.TfIdf()
    for desc in descriptions:
        tfidf.addDocument(desc)

    vectors = [tfidf.listTfIdf(desc) for desc in descriptions]

    kmeans = KMeans(num_clusters)
    clusters = kmeans.fit_predict(vectors)

    return clusters

class JobApplication:
    def __init__(self, user_id, cv_text, preferences):
        self.user_id = user_id
        self.cv_text = cv_text
        self.preferences = preferences

    def apply_to_jobs(self):
        """Apply to jobs based on CV analysis and job clustering"""
        try:
            cv_analyzer = CVAnalyzer(self.cv_text)
            tfidf = cv_analyzer.analyze_skills()

            jobs = self.fetch_jobs(self.preferences)  # Fixed method call

            clustered_jobs = cluster_jobs(jobs)

            response = self.apply_to_jobs_api(self.user_id, clustered_jobs)

            return response
        except Exception as e:
            print("Error applying to jobs:", e)
            raise

    def fetch_jobs(self, preferences):
        """Fetch jobs based on preferences."""
        # Replace this with actual job fetching logic
        jobs = [
            {"description": "Job description 1"},
            {"description": "Job description 2"},
            {"description": "Job description 3"},
            {"description": "Job description 4"},
            {"description": "Job description 5"},
        ]
        return jobs

    def apply_to_jobs_api(self, user_id, clustered_jobs):
        """Apply to Jobs using Anthropics API"""
        try:
            response = requests.post(  # Changed from axios to requests
                "https://anthropics-api.com/apply",
                json={"userId": user_id, "jobs": clustered_jobs},
                headers={"Authorization": f"Bearer {ANTHROPICS_API_KEY}", "Content-Type": "application/json"},
            )
            return response.json()
        except Exception as e:
            print("Error applying to jobs via API:", e)
            raise

class MailService:
    SMTP_SERVER = "smtp.gmail.com"
    PORT = 465
    SENDER_EMAIL = "your_email@gmail.com"
    PASSWORD = 'your_password_here'
    CONTEXT = ssl.create_default_context()

    def send_mail(self, template, receiver_email, name, subject):
        """Send email with plain text and HTML versions"""
        message = MIMEMultipart('alternative')
        message["Subject"] = subject
        message["From"] = self.SENDER_EMAIL
        message["To"] = receiver_email

        text = f"Hi, {name}\n{template}"
        html = f"""\
        <html>
          <body>
            <h2>Hi, {name}</h2>
            <p>{template}</p>
          </body>
        </html>
        """

        message.attach(MIMEText(text, "plain"))
        message.attach(MIMEText(html, "html"))

        try:
            with smtplib.SMTP_SSL(self.SMTP_SERVER, self.PORT, context=self.CONTEXT) as server:
                server.login(self.SENDER_EMAIL, self.PASSWORD)
                server.sendmail(self.SENDER_EMAIL, receiver_email, message.as_string())
        except Exception as e:
            print(f"Error sending email: {e}")


# Example usage:
if __name__ == "__main__":
    # Example CV text extraction
    cv_file_path = "path_to_your_cv.pdf"
    cv_text = extract_text(cv_file_path)

    # Example job preferences
    preferences = {"location": "Remote", "experience_level": "Senior"}

    # Example job application
    job_application = JobApplication("user123", cv_text, preferences)
    application_response = job_application.apply_to_jobs()
    print("Application response:", application_response)

    # Example email sending
    mail_service = MailService()
    mail_service.send_mail("This is a test email", "recipient@example.com", "John Doe", "Test Email")
